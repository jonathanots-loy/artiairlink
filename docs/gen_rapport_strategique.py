from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── PAGE MARGINS ──
section = doc.sections[0]
section.page_width  = Inches(8.27)   # A4
section.page_height = Inches(11.69)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.0)

# ── COLORS ──
G1  = RGBColor(0x23, 0x4d, 0x32)   # dark green
G2  = RGBColor(0x2f, 0x6b, 0x45)
G3  = RGBColor(0x4e, 0x8d, 0x65)
MUT = RGBColor(0x66, 0x75, 0x65)
RED = RGBColor(0xba, 0x5a, 0x4e)
AMB = RGBColor(0xd1, 0xa0, 0x5a)
INK = RGBColor(0x1a, 0x28, 0x1c)
WHT = RGBColor(0xff, 0xff, 0xff)

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = color
    return p

def add_para(doc, text, bold_parts=None, color=None, size=10, italic=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    run.bold = False
    run.italic = italic
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.4 + level * 0.4)
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, '234d32')
        run = cell.paragraphs[0].add_run(h)
        run.font.color.rgb = WHT
        run.font.bold = True
        run.font.size = Pt(9)
    # Data rows
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        bg = 'f5f2eb' if ri % 2 == 0 else 'fffdf8'
        for ci, cell_text in enumerate(row):
            cell = tr.cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            if isinstance(cell_text, tuple):
                run = p.add_run(cell_text[0])
                run.font.size = Pt(9)
                run.font.bold = cell_text[1]
                if len(cell_text) > 2:
                    run.font.color.rgb = cell_text[2]
            else:
                run = p.add_run(cell_text)
                run.font.size = Pt(9)
    # Column widths
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[ci].width = Cm(w)
    return table

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '8cbf6c')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ════════════════════════════════════════════════
# COVER
# ════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('AgroVision — ArtiAirLink')
run.font.size  = Pt(28)
run.font.bold  = True
run.font.color.rgb = G1

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('Conseil Stratégique : R&D · Vente · Juridique')
r2.font.size = Pt(15)
r2.font.color.rgb = G3

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run('Session du 18 Avril 2026  ·  Confidentiel')
r3.font.size = Pt(10)
r3.font.color.rgb = MUT
r3.italic = True

doc.add_paragraph()
add_divider(doc)
doc.add_paragraph()

# ════════════════════════════════════════════════
# 0. INTRO
# ════════════════════════════════════════════════
add_heading(doc, 'Votre réalité (sans filtre)', level=1, color=G1)
add_para(doc,
    'Vous êtes des ingénieurs industriels qui construisent un outil d\'analyse végétale sans drone réel, '
    'sans données de terrain validées, et sans expertise agronomique. C\'est exactement là qu\'ont commencé '
    'la plupart des grandes startups agritech — mais il faut nommer les lacunes pour les combler.',
    size=10)
add_para(doc,
    'Votre force : vous avez un produit fonctionnel (app PWA offline, multi-indices RGB, cartographie IDW, '
    'prescription maps). C\'est une base technique sérieuse que 95% des concurrents n\'ont pas au stade étudiant.',
    size=10)

add_divider(doc)

# ════════════════════════════════════════════════
# 1. R&D
# ════════════════════════════════════════════════
add_heading(doc, '1. R&D — Ce qui compte vraiment', level=1, color=G1)

add_heading(doc, 'Le problème fondamental à résoudre en priorité', level=2, color=G2)
add_para(doc,
    'Vos indices RGB (VARI, ExG, GRVI, GLI, TGI) sont des proxys optiques. Ils mesurent la couleur réfléchie, '
    'pas la santé réelle. Un plant sous stress hydrique et un plant sain peuvent avoir le même VARI si la lumière change. '
    'Sans données de calibration terrain corrélées à des analyses laboratoire réelles, votre app est un colorimètre amélioré.',
    size=10)

add_heading(doc, 'Partenaire académique prioritaire : Gembloux Agro-Bio Tech', level=2, color=G2)
add_para(doc,
    'Gembloux Agro-Bio Tech (Faculté de Sciences agronomiques de l\'ULiège, à 45 min de Bruxelles) est votre '
    'partenaire naturel. Ils disposent de parcelles expérimentales instrumentées, de chercheurs spécialisés en '
    'corrélation spectral/terrain, et d\'un historique de collaboration avec des startups.',
    size=10)
add_para(doc, 'Ce que vous leur proposez :', size=10)
add_bullet(doc, 'Vous fournissez : le logiciel + les vols drone')
add_bullet(doc, 'Eux fournissent : terrain expérimental, analyses labo, validation scientifique, co-publication')
add_para(doc,
    'Contact : Bureau des relations entreprises de Gembloux, ou repérez les publications du département '
    '"Biosystems Engineering" sur Google Scholar et contactez l\'auteur directement.',
    size=10, color=MUT, italic=True)

add_heading(doc, 'Pour la RDC — Le partenaire local est non-négociable', level=2, color=G2)
add_table(doc,
    ['Organisme', 'Rôle', 'Utilité pour vous'],
    [
        ('INERA (Kinshasa)', 'Recherche agronomique nationale', 'Validation cultures locales'),
        ('CGIAR / CIMMYT',   'Réseau recherche international (maïs, manioc)', 'Données maladies RDC'),
        ('UNIKIN — Agro',    'Université de Kinshasa', 'Partenaire académique local'),
    ],
    col_widths=[4.5, 5.0, 5.5]
)

doc.add_paragraph()
add_heading(doc, 'Roadmap R&D révisée (réaliste)', level=2, color=G2)
add_table(doc,
    ['Priorité', 'Action', 'Délai'],
    [
        ('🔴 Urgent',       'Partenariat Gembloux — 1 parcelle test, 1 culture',                                 'Maintenant'),
        ('🔴 Urgent',       'DJI Mini 3 Pro (~800€, <249g, légal Open A1 sans certificat)',                      'Dès que possible'),
        ('🟠 Court terme',  'Collecter 50+ captures corrélées à analyses chlorophylle/SPAD réelles',             'Printemps 2026'),
        ('🟠 Court terme',  'Phase 2B TF.js — entraîner sur données propres, pas seulement PlantVillage (US)',   'Été 2026'),
        ('🟡 Moyen terme',  'Maladies RDC : mosaïque manioc, chenille légionnaire du maïs, flétrissement banane','Fin 2026'),
    ],
    col_widths=[3.0, 9.0, 3.0]
)

doc.add_paragraph()
add_para(doc,
    'Pourquoi pas le DJI Mavic 3M maintenant ? Il coûte ~9 500€ et nécessite une autorisation Specific Category. '
    'Le Mini 3 Pro à 800€ vous donne déjà des données RGB exploitables pour valider votre modèle business '
    'avant d\'investir dans le multispectral.',
    size=10, color=MUT, italic=True)

add_divider(doc)

# ════════════════════════════════════════════════
# 2. VENTE
# ════════════════════════════════════════════════
add_heading(doc, '2. Vente — Ne vendez pas aux agriculteurs en premier', level=1, color=G1)

add_heading(doc, 'Erreur classique des startups agritech', level=2, color=G2)
add_para(doc,
    'Vendre directement aux agriculteurs individuels est lent, coûteux, et les cycles de vente durent 6-18 mois. '
    'Un agriculteur belge moyen a 45 ans, est pragmatique, et ne paiera que si vous lui prouvez un ROI '
    'concret en euros par hectare.',
    size=10)

add_heading(doc, 'Vos vraies cibles commerciales', level=2, color=G2)

add_para(doc, 'Niveau 1 — Accélérateurs de crédibilité (pas de revenu immédiat, mais ouvrent les portes)', size=10, color=G3)
add_bullet(doc, 'Concours startup : Enactus Belgium, Start it @KBC, Solvay Entrepreneurs, BeAngels')
add_bullet(doc, 'Incubateurs agri : AgriTech Flanders, Greenwin (cluster Wallonie), Wagralim')
add_bullet(doc, 'Programmes EU : EIC Accelerator (jusqu\'à 2,5M€ de grant + equity), Horizon Europe EIC Pathfinder')

doc.add_paragraph()
add_para(doc, 'Niveau 2 — Premiers revenus réels (B2B)', size=10, color=G3)
add_table(doc,
    ['Cible', 'Pays', 'Pourquoi ils paient', 'Prix indicatif'],
    [
        ('Coopératives agricoles (FUGEA, AgroFront)',    'Belgique',      'Service drone mutualisé pour leurs membres',               '5–15k€/an'),
        ('Assureurs agricoles (Axa Agri, Ethias)',       'Belgique',      'Données perte de récolte pour régler les sinistres',        'License data'),
        ('ONG agritech (WFP, SNV, IFAD)',                'RDC / Afrique', 'Monitoring récoltes à grande échelle',                     'Grant 20–100k€'),
        ('Fournisseurs d\'intrants (Syngenta, Bayer)',   'Belgique',      'Données précision pour recommander leurs produits',         'Partenariat'),
    ],
    col_widths=[4.5, 2.5, 5.5, 2.5]
)

doc.add_paragraph()
add_para(doc, 'Niveau 3 — SaaS à 3-5 ans', size=10, color=G3)
add_para(doc,
    '~15–25€/ha/an, vendu via coopératives comme intermédiaires. Un agriculteur wallon moyen cultive ~60ha → '
    '~1 200€/an par client. Une coopérative avec 100 membres → 120k€ ARR potentiel.',
    size=10)

add_heading(doc, 'Messages de vente par marché', level=2, color=G2)
add_table(doc,
    ['Marché', 'Message clé'],
    [
        ('Belgique', '"Réduisez vos intrants de 20–30% en traitant uniquement les zones stressées — '
                     'économies sur phytosanitaires, conformité directive nitrates, traçabilité labels bio."'),
        ('RDC',      '"Détectez la mosaïque du manioc 3 semaines avant les symptômes visibles — '
                     'évitez 40–60% de pertes de récolte sans coût en intrants chimiques."'),
    ],
    col_widths=[3.0, 12.0]
)

doc.add_paragraph()
add_heading(doc, 'Ce que vous devez construire côté vente', level=2, color=G2)
add_bullet(doc, '1 cas d\'usage documenté : 1 champ, 1 culture, avant/après, résultat chiffré en €')
add_bullet(doc, 'Formulaire "Demander un pilote gratuit" sur votre landing page site.html')
add_bullet(doc, 'Deck investisseur 10 slides : problème → solution → démo → traction → équipe → ask')

add_divider(doc)

# ════════════════════════════════════════════════
# 3. JURIDIQUE
# ════════════════════════════════════════════════
add_heading(doc, '3. Juridique — Ce que personne ne vous dit', level=1, color=G1)

add_heading(doc, 'Réglementation drones EASA (Belgique/EU)', level=2, color=G2)
add_table(doc,
    ['Drone', 'Poids', 'Catégorie', 'Requis'],
    [
        ('Drone jouet actuel',  '< 250g prob.',  'Open A1',          'Rien (usage non-commercial)'),
        ('DJI Mini 3 Pro',      '248g',          'Open A1',          'Enregistrement MySKYPortal DGTA si commercial'),
        ('DJI Mavic 3M',        '~895g',         'Specific Category','Autorisation DGTA + Certificat A2 minimum'),
    ],
    col_widths=[3.5, 2.5, 3.5, 5.5]
)
doc.add_paragraph()
add_para(doc,
    'Pour la RDC : l\'Autorité de l\'Aviation Civile du Congo (AAC) régule les drones commerciaux. '
    'La réglementation est moins formalisée qu\'en EU — un partenaire local est indispensable pour naviguer les autorisations.',
    size=10, color=MUT, italic=True)

add_heading(doc, 'RGPD et données agricoles', level=2, color=G2)
add_para(doc,
    'Bonne nouvelle : votre app actuelle (100% locale, IndexedDB, aucun serveur) n\'est PAS soumise au RGPD. '
    'Vous ne collectez aucune donnée personnelle et ne les transmettez à aucun serveur.',
    size=10)
add_para(doc, 'Dès que vous ajoutez un backend (Phase 4) :', size=10)
add_bullet(doc, 'Les coordonnées GPS d\'un champ peuvent identifier un propriétaire → données personnelles potentielles')
add_bullet(doc, 'Obligation : politique de confidentialité, mentions légales, DPA (Data Processing Agreement) avec clients')
add_bullet(doc, 'Hébergez en Europe (Hetzner, OVH) pour simplifier la conformité RGPD')

add_heading(doc, 'Propriété intellectuelle', level=2, color=G2)
add_para(doc, 'Ce que vous pouvez protéger maintenant (gratuitement) :', size=10)
add_bullet(doc, 'Copyright : votre code est automatiquement protégé dès sa création. Ajoutez © 2026 ArtiAirLink dans vos fichiers.')
add_bullet(doc, 'Marque "AgroVision" : vérifiez la disponibilité sur boip.int (BOIP). Dépôt Benelux ≈ 350€ pour 10 ans.')

add_para(doc, 'Ce que vous ne devriez PAS faire maintenant :', size=10)
add_bullet(doc, 'Brevet : coûte 5–15k€ minimum, 2–4 ans de procédure. Inutile à ce stade — misez sur le first-mover advantage.')

add_heading(doc, 'Structure juridique — Créez la société au bon moment', level=2, color=G2)
add_para(doc, 'Ne créez PAS de SRL tout de suite. Attendez d\'avoir :', size=10)
add_bullet(doc, 'Un co-fondateur clairement identifié + accord écrit sur la répartition des parts')
add_bullet(doc, 'Au moins 1 client payant ou 1 subvention confirmée')
add_bullet(doc, 'Clarté sur qui travaille à temps plein vs temps partiel')

add_para(doc, 'En attendant, utilisez le statut Étudiant-Entrepreneur belge :', size=10)
add_bullet(doc, 'Disponible dans la plupart des universités belges')
add_bullet(doc, 'Permet de facturer légalement sans créer de société')
add_bullet(doc, 'Parfois accompagné d\'une bourse mensuelle de 500–700€')

add_para(doc, 'Quand vous créerez la SRL :', size=10)
add_bullet(doc, 'Capital minimum : 1 euro (Code des Sociétés belge 2019)')
add_bullet(doc, 'Objet social large : "développement, exploitation et commercialisation de logiciels et services de drone à usage agricole"')
add_bullet(doc, 'Pacte d\'actionnaires entre co-fondateurs AVANT la création (clauses de vesting, drag-along, first refusal) — budget 500–1 500€ chez un avocat startup belge')

add_divider(doc)

# ════════════════════════════════════════════════
# 4. ACTIONS CONCRÈTES
# ════════════════════════════════════════════════
add_heading(doc, '4. Les 5 actions à faire cette semaine', level=1, color=G1)

add_table(doc,
    ['#', 'Action', 'Impact'],
    [
        ('1', 'Envoyer un email à Gembloux Agro-Bio Tech (5 lignes, proposez un appel de 20 min)',
              'Valide scientifiquement votre produit'),
        ('2', 'Vérifier la disponibilité de "AgroVision" sur boip.int',
              'Protège votre marque avant d\'avoir des clients'),
        ('3', 'S\'inscrire sur MySKYPortal DGTA (gratuit, 30 min)',
              'Crédibilité + conformité légale vol commercial'),
        ('4', 'Identifier 3 agriculteurs/coopératives pour un pilote gratuit ce printemps',
              '1 pilote documenté > 6 mois de dev pour vendre'),
        ('5', 'Ajouter formulaire "Rejoindre le pilote" sur site.html',
              'Convertit votre trafic GitHub/Netlify en leads'),
    ],
    col_widths=[0.8, 9.7, 4.5]
)

doc.add_paragraph()

# ════════════════════════════════════════════════
# CONCLUSION
# ════════════════════════════════════════════════
add_heading(doc, 'Conclusion', level=1, color=G1)
p = doc.add_paragraph()
r = p.add_run(
    'Le vrai avantage compétitif d\'AgroVision n\'est pas technique — c\'est que vous construisez pour des marchés '
    '(RDC + Belgique) que les grands acteurs (John Deere, Trimble, DJI AgriScan) considèrent trop petits ou trop risqués. '
    'Cette combinaison est votre créneau réel. Capitalisez dessus avant que quelqu\'un d\'autre ne le fasse.'
)
r.font.size = Pt(10)
r.italic = True
r.font.color.rgb = G1

add_divider(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('ArtiAirLink · AgroVision · contact : jonathan.ots1@gmail.com')
r.font.size = Pt(9)
r.font.color.rgb = MUT

# ── SAVE ──
out = r'C:\Users\jonat\OneDrive\Bureau\ARTI\artiairlink\docs\Rapport_Strategique_AgroVision_20260418.docx'
doc.save(out)
print('Saved:', out)
